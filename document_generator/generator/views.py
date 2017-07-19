import random
import string

from django.shortcuts import render_to_response, render, redirect
from .models import *
from .forms import *
from .custom_decorators import is_deanery, is_student
from django.contrib.auth import get_user_model



@is_deanery
def ind_list(request):
    edu = IndividualTask.objects.filter(practice_type=IndividualTask.EDU)
    prod = IndividualTask.objects.filter(practice_type=IndividualTask.PROD)
    dip = IndividualTask.objects.filter(practice_type=IndividualTask.DIP)
    return render(request, 'deanery/ind_tasks.html', locals())


@is_deanery
def ind_edit(request, id):
    title = "Редактирование индивидуального задания"
    ind = IndividualTask.objects.get(pk=id)
    if request.POST:
        form = IndividualTaskForm(request.POST or None, instance=ind)
        if form.is_valid():
            form.save()
            return redirect("/ind_tasks/")
        else:
            return render(request, 'deanery/ind_task.html', locals())
    else:
        form = IndividualTaskForm(instance=ind)
        return render(request, 'deanery/ind_task.html', locals())


@is_deanery
def new_ind(request, type):
    if type == 'edu':
        choice = IndividualTask.EDU
    elif type == 'prod':
        choice = IndividualTask.PROD
    elif type == 'dip':
        choice = IndividualTask.DIP
    else:
        return redirect('/ind_tasks/')
    if request.POST:
        number = IndividualTask.objects.filter(practice_type=choice).count() + 1
        form = IndividualTaskForm(request.POST or None, initial={
            'practice_type': choice,
            'task_number': number
        })
        if form.is_valid():
            form.save()
            return redirect('/ind_tasks/')
        else:
            return render(request, 'deanery/addIndTask.html', {'form': form})
    else:
        form = IndividualTaskForm(initial={
            'practice_type': choice
        })
        return render(request, 'deanery/addIndTask.html', {'form': form})


@is_deanery
def new_practice(request):
    title = "Добавление практики"
    if request.POST:
        form = PracticeForm(request.POST or None)
        if form.is_valid():
            form.save()
            return redirect("/practices/")
        else:
            return render(request, 'deanery/addPractice.html', {'form': form})
    else:
        form = PracticeForm()
        return render(request, 'deanery/addPractice.html', {'form': form})


@is_deanery
def edit_practice(request, id):
    title = "Редактирование практики"
    practice = Practice.objects.get(pk=id)
    # students_list = Student.objects.get(practice=practice)
    students_list = Student.objects.all().filter(practice=practice)
    if request.POST:
        form = PracticeForm(request.POST or None, instance=practice)
        if form.is_valid():
            form.save()
            return redirect("/practices/%s/" % id)
        else:
            return render(request, 'deanery/practice.html', {'form': form})
    else:
        form = PracticeForm(instance=practice)
        return render(request, 'deanery/practice.html', locals())


@is_deanery
def students(request):
    title = "Студенты"
    students_list = Student.objects.all()
    return render(request, 'deanery/students.html', locals())


@is_deanery
def new_student(request):
    title = "Добавление студента"
    if request.POST:
        form = StudentForm(request.POST or None)
        if form.is_valid():
            student = form.save(commit=False)
            login = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
            password = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
            user = get_user_model().objects.create_user(username=login, password=password)
            Group.objects.get(name='Students').user_set.add(user)
            student.login = login
            student.password = password
            student.user = user
            student.save()
            form.save_m2m()
            # creating docs
            return redirect("/students/")
        else:
            return render(request, 'deanery/addStudent.html', {'form': form})
    else:
        form = StudentForm()
        return render(request, 'deanery/addStudent.html', {'form': form})


@is_deanery
def edit_student(request, id):
    title = "Редактирование студента"
    student = Student.objects.get(pk=id)
    if request.POST:
        form = StudentForm(request.POST or None, instance=student)
        if form.is_valid():
            form.save()
            return redirect("/students/")
    else:
        form = StudentForm(instance=student)
        return render(request, 'deanery/student.html', {'form': form}, locals())


@is_deanery
def practices(request):
    title = "Практики"
    practices_list = Practice.objects.all()
    return render(request, 'deanery/practices.html', locals())


@is_student
def profile(request):
    student = Student.objects.get(user=request.user)
    if request.POST:
        form = StudentForm(request.POST or None, instance=student)
        if form.is_valid():
            form.save()
            return redirect("/")
        else:
            return render(request, 'student/profile.html', locals())
    else:
        form = StudentForm(instance=student)
        return render(request, 'student/profile.html', locals())


@is_student
def practice_docs(request, id):
    return render(request, 'student/docs.html', locals())


# TODO FINISH THAT
@is_student
def edit_individual(request, id):
    practice = Practice.objects.filter(pk=id)
    tasks = IndividualTask.objects.filter(practice_type=practice.practice_type)
    if request.POST:
        pass
    pass


@is_student
def edit_diary(request):
    pass


@is_student
def edit_pass(request, id):
    pass


@is_student
def pass_view(request, id):
    pass


@is_student
def new_diary_record(request, id):
    practice = Practice.objects.get(pk=id)
    d = Diary.objects.get(practice=practice)
    if request.POST:
        form = DiaryRecordForm(request, initial={'diary': d})
        if form.is_valid():
            form.save()
            return redirect('/practice_%s/diary/' % id)
        else:
            return render(request, 'student/addDiaryRecord.html', {'form': form})
    else:
        form = DiaryRecordForm()
        return render(request, 'student/addDiaryRecord.html', {'form': form})


@is_student
def diary_view(request, id):
    pass


@is_student
def diary(request, id):
    practice = Practice.objects.get(pk=id)
    diary = Diary.objects.get(practice=practice)
    records = DiaryRecord.objects.filter(diary=diary).order_by('date')
    return render(request, 'student/diary.html', locals())


@is_student
def edit_record(request, id, record_id):
    practice = Practice.objects.get(pk=id)
    diary = Diary.objects.get(practice=practice)
    record = DiaryRecord.objects.get(pk=record_id)
    if request.POST:
        form = DiaryRecordForm(request.POST or None, instance=record)
        if form.is_valid():
            form.save()
            return redirect('/practice_%s/diary' % id)
        else:
            return render(request, 'student/record.html', {'form': form})
    else:
        form = DiaryRecordForm(instance=record)
        return render(request, 'student/record.html', locals())


@is_student
def diary_download(request, id):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    diary = Document()
    font = diary.styles['Normal'].font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    header = diary.add_paragraph(
        'Дневник студента-практиканта\nХод выполнения практики\n').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = diary.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    # Создание и заполнение шапки
    hat = table.rows[0]
    hat.cells[0].add_paragraph('№').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hat.cells[1].add_paragraph('Описание выполненной работы').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hat.cells[2].add_paragraph('Дата').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hat.cells[3].add_paragraph('Отметка\nруководителя').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # data
    practice = Practice.objects.get(pk=id)
    diary = Diary.objects.get(practice=practice)
    records = DiaryRecord.objecst.filter(diary=diary)


    # Заполнение таблицы
    for i in range(10):
        nextRow = table.add_row()
        nextRow.cells[0].add_paragraph(str(i + 1)).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nextRow.cells[1].add_paragraph('some_text').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nextRow.cells[2].add_paragraph('some_date').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nextRow.cells[3].add_paragraph('some_mark').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Форматирование ширины таблицы
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(0.35)
    for cell in table.columns[1].cells:
        cell.width = Inches(3)
    for cell in table.columns[2].cells:
        cell.width = Inches(1.65)
    for cell in table.columns[3].cells:
        cell.width = Inches(1.65)

    diary.save('prepairDocx/Заполненный дневник.docx')


@is_student
def report_view(request, id):
    pass


@is_student
def report_download(request, id):
    pass


@is_student
def individual_view(request, id):
    pass


@is_student
def individual_download(request, id):
    pass


@is_student
def pass_download(request, id):
    pass
